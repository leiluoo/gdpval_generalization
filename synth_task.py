"""
Pipeline 2: Synthesize task prompts and rubrics from Pool C files.

For each Pool C reference file, reads its content, looks at the original
task's occupation/sector/task_type, and generates a new (prompt, rubric)
pair that references the new file's actual entities.

Usage:
    python synth_task.py
    python synth_task.py --limit 5
"""

import json
import argparse
from pathlib import Path

from llm_client import call_claude, extract_json

OUTPUT_DIR = Path("data/pool_c")
TASKS_PATH = Path("data/single_ref_tasks.json")


# ── file readers ──────────────────────────────────────────────────────────────

def read_xlsx_summary(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for sname in wb.sheetnames:
        ws = wb[sname]
        lines.append(f"[Sheet: {sname}] {ws.max_row} rows x {ws.max_column} cols")
        for i, row in enumerate(ws.iter_rows(max_row=5, values_only=True)):
            vals = [str(v)[:30] if v is not None else "" for v in row]
            lines.append(f"  row{i+1}: {vals}")
    return "\n".join(lines)


def read_docx_summary(path: str) -> str:
    from docx import Document
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(f"[{p.style.name}] {p.text[:120]}")
    for i, t in enumerate(doc.tables):
        lines.append(f"[Table {i+1}: {len(t.rows)}x{len(t.columns)}]")
        for ri, row in enumerate(t.rows[:4]):
            lines.append(f"  row{ri}: {[c.text[:30] for c in row.cells[:5]]}")
    return "\n".join(lines)


def read_file_summary(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".xlsx":
        return read_xlsx_summary(path)
    elif ext == ".docx":
        return read_docx_summary(path)
    return f"[{ext} file, content not parsed]"


# ── task type extraction ──────────────────────────────────────────────────────

def extract_task_type(original_task: dict) -> str:
    """Extract task pattern from rubric — describes output, not input data."""
    rubric_lines = original_task["rubric_pretty"].split("\n")[:12]
    deliv_types = list({f.split(".")[-1].lower() for f in original_task["deliverable_files"]})
    first_sentence = original_task["prompt"].split(".")[0][:200]
    return (
        f"Occupation: {original_task['occupation']} | Sector: {original_task['sector']}\n"
        f"Deliverable types: {deliv_types}\n"
        f"Task summary: {first_sentence}\n"
        f"Rubric excerpt:\n" + "\n".join(rubric_lines)
    )


# ── prompt + rubric synthesis ─────────────────────────────────────────────────

def synthesize_task(pool_c_meta: dict, original_task: dict) -> dict:
    file_path = pool_c_meta["pool_c_file"]
    file_summary = read_file_summary(file_path)
    task_type_desc = extract_task_type(original_task)
    deliv_types = list({f.split(".")[-1].lower() for f in original_task["deliverable_files"]})
    v = pool_c_meta["variant_idx"]

    prompt = f"""You are creating a synthetic professional task for AI training data.

== ORIGINAL TASK PATTERN ==
{task_type_desc}

== NEW REFERENCE FILE ==
Filename: {Path(file_path).name}
{file_summary}

== YOUR JOB ==
Write a complete task specification grounded in the NEW file's content above.

Requirements:
1. Professional role matching the occupation
2. Realistic business scenario using entities FROM the new file (column headers, sheet names, sample values visible above)
3. Same TYPE of deliverable as original: {deliv_types}
4. Concrete, objectively-evaluable requirements (avoid vague language)
5. 10-14 rubric criteria covering format, content, and computation correctness

Return ONLY this JSON (no extra text):
{{
  "prompt": "Full multi-paragraph task prompt...",
  "rubric_pretty": "[+2] Criterion 1\\n[+2] Criterion 2\\n...",
  "rubric_json": [
    {{"score": 2, "criterion": "Exact criterion text", "rubric_item_id": "syn-{v}-1"}},
    {{"score": 2, "criterion": "...", "rubric_item_id": "syn-{v}-2"}}
  ],
  "deliverable_files": ["output_name.xlsx"],
  "task_summary": "One-sentence description"
}}
"""

    spec = extract_json(call_claude(prompt))

    return {
        "original_task_id": original_task["task_id"],
        "occupation": original_task["occupation"],
        "sector": original_task["sector"],
        "pool_c_file": file_path,
        "file_type": pool_c_meta["file_type"],
        "variant_idx": v,
        "prompt": spec.get("prompt", ""),
        "rubric_pretty": spec.get("rubric_pretty", ""),
        "rubric_json": spec.get("rubric_json", []),
        "deliverable_files": spec.get("deliverable_files", []),
        "task_summary": spec.get("task_summary", ""),
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--pool_c_meta", default="data/pool_c/pool_c_metadata.json")
    parser.add_argument("--limit", type=int, default=None)
    args = parser.parse_args()

    with open(TASKS_PATH) as f:
        all_tasks = json.load(f)
    task_by_id = {t["task_id"]: t for t in all_tasks}

    with open(args.pool_c_meta) as f:
        pool_c = json.load(f)

    results_path = OUTPUT_DIR / "synthesized_tasks.json"
    results = json.loads(results_path.read_text()) if results_path.exists() else []
    done_keys = {(r["original_task_id"], r["variant_idx"]) for r in results}

    candidates = pool_c[:args.limit] if args.limit else pool_c

    for meta in candidates:
        key = (meta["original_task_id"], meta["variant_idx"])
        if key in done_keys:
            print(f"  Skip: {meta['original_task_id'][:8]} v{meta['variant_idx']}")
            continue

        original = task_by_id.get(meta["original_task_id"])
        if not original:
            continue

        print(f"  Task synthesis: {meta['original_task_id'][:8]} v{meta['variant_idx']} [{meta['file_type']}] ...")
        try:
            task_spec = synthesize_task(meta, original)
            results.append(task_spec)
            results_path.write_text(json.dumps(results, indent=2))
            print(f"    prompt: {len(task_spec['prompt'])} chars, rubric: {len(task_spec['rubric_json'])} items")
        except Exception as e:
            print(f"    ERROR: {e}")

    print(f"\nDone. {len(results)} synthesized tasks saved.")


if __name__ == "__main__":
    main()
