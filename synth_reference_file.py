"""
Pipeline 1: Synthesize reference files from schema.

Each variant is assigned a concrete diversity spec (industry × region × scale × currency)
so Claude generates genuinely different scenarios, not just re-rolls of the same prompt.

Usage:
    python synth_reference_file.py --task_id <id> --variants 5
    python synth_reference_file.py --all --variants 5
"""

import json
import os
import random
import argparse
import requests
import re
from pathlib import Path

from llm_client import call_claude, extract_json
from extract_schema import extract_schema
from diversity_spec import build_diversity_spec, diversity_instruction


OUTPUT_DIR = Path("data/pool_c")
REF_CACHE_DIR = Path("data/ref_cache")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
REF_CACHE_DIR.mkdir(parents=True, exist_ok=True)


# ── helpers ───────────────────────────────────────────────────────────────────

def download_ref_file(task: dict) -> Path:
    url = task["reference_file_urls"][0]
    fname = task["reference_files"][0].split("/")[-1]
    safe = re.sub(r"[^\w.\-]", "_", fname)
    dest = REF_CACHE_DIR / f"{task['task_id'][:8]}_{safe}"
    if not dest.exists():
        r = requests.get(url, timeout=60)
        r.raise_for_status()
        dest.write_bytes(r.content)
    return dest


# ── XLSX synthesis ────────────────────────────────────────────────────────────

def synth_xlsx(task: dict, schema: dict, variant_idx: int, div_spec: dict) -> Path:
    from openpyxl import Workbook

    base_row_count = max(
        s.get("row_count", 100) for s in schema.get("sheets", [{"row_count": 100}])
    )
    # Row count is now derived from computation complexity, not a raw scale factor
    computation = div_spec.get("computation", "")
    row_count = base_row_count
    if "multi-step" in computation or "conditional" in computation:
        row_count = max(50, base_row_count // 2)   # complex formulas → smaller dataset
    elif "raw data" in computation:
        row_count = min(base_row_count * 2, 3000)  # raw data → can go larger

    prompt = f"""You are generating a synthetic Excel dataset for AI training purposes.

OCCUPATION: {task['occupation']}
SECTOR: {task['sector']}
ORIGINAL FILE STRUCTURE (schema only — no raw data):
{json.dumps(schema, indent=2)}

{diversity_instruction(div_spec, "xlsx")}

TASK: Design a new Excel file that satisfies the structural constraints above.
The file must be genuinely different from the original in structure (not just renamed columns).
- Column headers and metric names must fit the domain specified above
- All entity names (companies, divisions, countries) must be completely fictional
- Numeric ranges must be realistic for the domain
- If the data_structure constraint requires multi-sheet or cross-tab, implement it

Return ONLY a JSON object:
{{
  "scenario_description": "One sentence describing the fictional organization and what this file tracks",
  "sheets": [
    {{
      "name": "SheetName",
      "columns": [
        {{
          "header": "Column Name",
          "type": "categorical|numeric|integer|date|text",
          "values": ["val1", "val2", ...],
          "range": [min, max]
        }}
      ],
      "row_count": {row_count}
    }}
  ]
}}
"""

    spec = extract_json(call_claude(prompt))

    wb = Workbook()
    wb.remove(wb.active)

    for sheet_spec in spec["sheets"]:
        ws = wb.create_sheet(sheet_spec["name"])
        cols = sheet_spec["columns"]
        row_count = sheet_spec.get("row_count", row_count)

        for ci, col in enumerate(cols, 1):
            ws.cell(row=1, column=ci, value=col["header"])

        col_pools = {}
        for col in cols:
            ctype = col["type"]
            if ctype == "categorical":
                col_pools[col["header"]] = col.get("values", ["A", "B", "C"])
            elif ctype in ("numeric", "integer"):
                lo, hi = col.get("range", [0, 1000])
                col_pools[col["header"]] = (lo, hi, ctype)
            elif ctype == "date":
                col_pools[col["header"]] = ("date",)
            else:
                col_pools[col["header"]] = col.get("values", ["Value"])

        for row_idx in range(2, row_count + 2):
            for ci, col in enumerate(cols, 1):
                h = col["header"]
                pool = col_pools[h]
                if isinstance(pool, tuple) and pool[0] == "date":
                    from datetime import date, timedelta
                    val = (date(2023, 1, 1) + timedelta(days=random.randint(0, 730))).isoformat()
                elif isinstance(pool, tuple) and len(pool) == 3:
                    lo, hi, ctype2 = pool
                    val = random.randint(int(lo), int(hi)) if ctype2 == "integer" else round(random.uniform(lo, hi), 2)
                elif isinstance(pool, list):
                    val = random.choice(pool)
                else:
                    val = ""
                ws.cell(row=row_idx, column=ci, value=val)

    out_path = OUTPUT_DIR / f"{task['task_id'][:8]}_xlsx_v{variant_idx}.xlsx"
    wb.save(out_path)
    return out_path


# ── DOCX synthesis ────────────────────────────────────────────────────────────

def synth_docx(task: dict, schema: dict, variant_idx: int, div_spec: dict) -> Path:
    from docx import Document

    prompt = f"""You are generating a synthetic Word document for AI training purposes.

OCCUPATION: {task['occupation']}
SECTOR: {task['sector']}
ORIGINAL DOCUMENT STRUCTURE (schema only — no raw content):
{json.dumps(schema, indent=2)}

{diversity_instruction(div_spec, "docx")}

TASK: Write the full content of a new Word document that satisfies the constraints above.
The document type and structural complexity must match the constraints — do not default to
the simplest possible structure.
- Paragraph style sequence should be similar to the original (same heading levels, list styles)
- Table row/col counts should match the original
- All entities, product names, companies, and numbers must be completely fictional
- Internally consistent and realistic for the specified domain

Return ONLY JSON:
{{
  "scenario_description": "One sentence describing the document",
  "paragraphs": [
    {{"style": "Normal", "text": "..."}}
  ],
  "tables": [
    {{
      "rows": [
        ["Cell 1", "Cell 2", "Cell 3"],
        ["Data", "1234.00", "unit"]
      ]
    }}
  ]
}}
"""

    spec = extract_json(call_claude(prompt))

    doc = Document()
    for para_spec in spec.get("paragraphs", []):
        p = doc.add_paragraph(para_spec.get("text", ""))
        try:
            p.style = doc.styles[para_spec.get("style", "Normal")]
        except KeyError:
            pass

    for table_spec in spec.get("tables", []):
        rows_data = table_spec.get("rows", [])
        if not rows_data:
            continue
        n_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=n_cols)
        for ri, row_data in enumerate(rows_data):
            for ci, cell_val in enumerate(row_data):
                table.rows[ri].cells[ci].text = str(cell_val)
        doc.add_paragraph()

    out_path = OUTPUT_DIR / f"{task['task_id'][:8]}_docx_v{variant_idx}.docx"
    doc.save(out_path)
    return out_path


# ── PDF synthesis ─────────────────────────────────────────────────────────────

def synth_pdf(task: dict, schema: dict, variant_idx: int, div_spec: dict) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    prompt = f"""You are generating a synthetic PDF document for AI training purposes.

OCCUPATION: {task['occupation']}
SECTOR: {task['sector']}
ORIGINAL DOCUMENT STRUCTURE (schema only):
{json.dumps(schema, indent=2)}

{diversity_instruction(div_spec, "pdf")}

TASK: Generate the full content of a document that satisfies the constraints above.
- The doc_type constraint defines the genre — follow it strictly
- The content_density constraint defines the ratio of tables to narrative
- All organizations, names, and data must be completely fictional

Return ONLY JSON:
{{
  "title": "Document Title",
  "sections": [
    {{
      "heading": "Section Heading",
      "content": "Paragraph text...",
      "table": [["Col1", "Col2"], ["r1c1", "r1c2"]]
    }}
  ]
}}
"""

    spec = extract_json(call_claude(prompt))
    out_path = OUTPUT_DIR / f"{task['task_id'][:8]}_pdf_v{variant_idx}.pdf"

    styles = getSampleStyleSheet()
    doc_rl = SimpleDocTemplate(str(out_path), pagesize=letter)
    story = [Paragraph(spec.get("title", "Document"), styles["Title"]), Spacer(1, 12)]

    for sec in spec.get("sections", []):
        if sec.get("heading"):
            story.append(Paragraph(sec["heading"], styles["Heading2"]))
        if sec.get("content"):
            story.append(Paragraph(sec["content"], styles["Normal"]))
        if sec.get("table"):
            tbl = Table(sec["table"])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            story.append(tbl)
        story.append(Spacer(1, 8))

    doc_rl.build(story)
    return out_path


# ── dispatcher ────────────────────────────────────────────────────────────────

def synthesize_reference_file(task: dict, variant_idx: int = 0) -> dict:
    ref_path = download_ref_file(task)
    schema = extract_schema(str(ref_path))
    ext = schema["file_type"]

    task_seed = hash(task["task_id"]) & 0xFFFF
    div_spec = build_diversity_spec(ext, variant_idx, task_seed)
    axes_summary = " | ".join(f"{k}={v.split('—')[0].strip()[:30]}" for k, v in div_spec.items())
    print(f"  [{ext}] {task['task_id'][:8]} v{variant_idx} | {axes_summary}")

    if ext == "xlsx":
        out_path = synth_xlsx(task, schema, variant_idx, div_spec)
    elif ext == "docx":
        out_path = synth_docx(task, schema, variant_idx, div_spec)
    elif ext == "pdf":
        out_path = synth_pdf(task, schema, variant_idx, div_spec)
    else:
        print(f"  Skipping unsupported type: {ext}")
        return {}

    return {
        "original_task_id": task["task_id"],
        "occupation": task["occupation"],
        "sector": task["sector"],
        "original_ref_file": task["reference_files"][0],
        "file_type": ext,
        "variant_idx": variant_idx,
        "diversity_spec": div_spec,
        "pool_c_file": str(out_path),
        "schema": schema,
    }


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--task_id", help="Task ID prefix to process")
    parser.add_argument("--all", action="store_true", help="Process all eligible tasks")
    parser.add_argument("--variants", type=int, default=3, help="Variants per task")
    parser.add_argument("--file_types", default="xlsx,docx,pdf")
    args = parser.parse_args()

    with open("data/single_ref_tasks.json") as f:
        tasks = json.load(f)

    allowed_types = set(args.file_types.split(","))
    eligible = [t for t in tasks if t["reference_files"][0].rsplit(".", 1)[-1].lower() in allowed_types]

    if args.task_id:
        eligible = [t for t in eligible if t["task_id"].startswith(args.task_id)]
    elif not args.all:
        eligible = eligible[:2]

    results_path = OUTPUT_DIR / "pool_c_metadata.json"
    results = json.loads(results_path.read_text()) if results_path.exists() else []
    done_keys = {(r["original_task_id"], r["variant_idx"]) for r in results}

    for task in eligible:
        for v in range(args.variants):
            if (task["task_id"], v) in done_keys:
                print(f"  Skip: {task['task_id'][:8]} v{v}")
                continue
            try:
                meta = synthesize_reference_file(task, variant_idx=v)
                if meta:
                    results.append(meta)
                    results_path.write_text(json.dumps(results, indent=2))
                    print(f"  -> {meta['pool_c_file']}")
            except Exception as e:
                print(f"  ERROR {task['task_id'][:8]} v{v}: {e}")

    print(f"\nDone. {len(results)} Pool C files recorded.")


if __name__ == "__main__":
    main()
