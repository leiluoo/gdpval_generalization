"""
Extract structural schema from reference files (xlsx/docx/pdf).
Produces a schema JSON that captures structure without leaking content.
"""

import json
import re
from pathlib import Path
from collections import defaultdict


def extract_xlsx_schema(file_path: str) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=False)
    sheets = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Find actual data bounds (skip empty rows/cols at edges)
        all_values = list(ws.iter_rows(values_only=True))
        non_empty_rows = [r for r in all_values if any(v is not None for v in r)]
        if not non_empty_rows:
            sheets.append({"name": sheet_name, "empty": True})
            continue

        max_col = max(len(r) for r in non_empty_rows)
        row_count = len(non_empty_rows)

        # Detect header row (first non-empty row with mostly string values)
        header_row = None
        header_idx = 0
        for i, row in enumerate(non_empty_rows[:5]):
            non_null = [v for v in row if v is not None]
            if non_null and sum(isinstance(v, str) for v in non_null) >= len(non_null) * 0.5:
                header_row = [str(v) if v is not None else None for v in row]
                header_idx = i
                break

        # Infer column types from data rows
        col_types = defaultdict(set)
        col_sample_values = defaultdict(list)
        for row in non_empty_rows[header_idx + 1: header_idx + 20]:
            for ci, val in enumerate(row):
                if val is not None:
                    col_types[ci].add(type(val).__name__)
                    if len(col_sample_values[ci]) < 3:
                        col_sample_values[ci].append(val)

        columns = []
        for ci in range(max_col):
            col_info = {
                "index": ci,
                "header": header_row[ci] if header_row and ci < len(header_row) else None,
                "types": list(col_types[ci]) if col_types[ci] else ["unknown"],
            }
            # For numeric columns, capture value range
            if col_types[ci] == {"int"} or col_types[ci] == {"float"} or col_types[ci] == {"int", "float"}:
                nums = [v for v in col_sample_values[ci] if isinstance(v, (int, float))]
                if nums:
                    col_info["sample_range"] = [min(nums), max(nums)]
            # For string columns, sample a few unique values (anonymized as examples)
            elif "str" in col_types[ci]:
                col_info["sample_values_count"] = len(set(str(v) for v in col_sample_values[ci]))
            columns.append(col_info)

        # Extract formulas (just the formula patterns, not values)
        formulas = []
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    # Extract formula pattern without cell references
                    formula_type = re.sub(r'[A-Z]+\d+', 'REF', cell.value)
                    if formula_type not in formulas:
                        formulas.append(formula_type)
                    if len(formulas) >= 10:
                        break
            if len(formulas) >= 10:
                break

        sheets.append({
            "name": sheet_name,
            "row_count": row_count,
            "col_count": max_col,
            "has_header": header_row is not None,
            "columns": columns,
            "formula_patterns": formulas[:10],
        })

    return {
        "file_type": "xlsx",
        "sheet_count": len(wb.sheetnames),
        "sheets": sheets,
    }


def extract_docx_schema(file_path: str) -> dict:
    from docx import Document

    doc = Document(file_path)

    # Extract paragraph structure (styles only, not content)
    para_structure = []
    for p in doc.paragraphs:
        if p.text.strip():
            para_structure.append({
                "style": p.style.name,
                "word_count": len(p.text.split()),
                "has_numbers": bool(re.search(r'\d[\d,\.]+', p.text)),
            })

    # Extract table structure
    tables = []
    for t in doc.tables:
        tables.append({
            "rows": len(t.rows),
            "cols": len(t.columns),
            "has_header": True,  # assume first row is header
            "col_headers": [c.text.strip()[:50] for c in t.rows[0].cells] if t.rows else [],
            "sample_row_types": _infer_table_col_types(t),
        })

    return {
        "file_type": "docx",
        "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
        "table_count": len(doc.tables),
        "paragraph_structure": para_structure,
        "tables": tables,
    }


def _infer_table_col_types(table) -> list:
    col_types = []
    if len(table.rows) < 2:
        return col_types
    for ci in range(len(table.columns)):
        values = [row.cells[ci].text.strip() for row in table.rows[1:] if ci < len(row.cells)]
        has_numbers = any(re.search(r'[\d,\.]+', v) for v in values)
        col_types.append("numeric" if has_numbers else "text")
    return col_types


def extract_pdf_schema(file_path: str) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        lines = [l.strip() for l in text.split("\n") if l.strip()]

        # Detect tables by looking for repeated delimiter patterns
        table_like_lines = [l for l in lines if re.search(r'\s{3,}', l)]

        pages.append({
            "page": i + 1,
            "line_count": len(lines),
            "table_like_line_count": len(table_like_lines),
            "has_dates": bool(re.search(r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b', text)),
            "has_numbers": bool(re.search(r'\$[\d,]+|\d+%', text)),
        })

    return {
        "file_type": "pdf",
        "page_count": len(reader.pages),
        "pages": pages,
    }


def extract_schema(file_path: str) -> dict:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".xlsx":
        return extract_xlsx_schema(file_path)
    elif ext == ".docx":
        return extract_docx_schema(file_path)
    elif ext == ".pdf":
        return extract_pdf_schema(file_path)
    else:
        return {"file_type": ext, "error": "unsupported type"}


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python extract_schema.py <file_path>")
        sys.exit(1)
    schema = extract_schema(sys.argv[1])
    print(json.dumps(schema, indent=2, default=str))
